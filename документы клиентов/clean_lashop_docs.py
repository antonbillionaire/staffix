"""
Очистка экспортов Telegram-канала LASHop от системного мусора:

Удаляются:
  - Заголовки канала, "Photo", "Video file", "Not included...", размеры файлов
  - Даты и время отдельными строками
  - Реакции ("❤4 👍2")
  - Ссылки t.me, instagram.com, http(s)://...
  - Хэштеги-слова (#ресницы), но НЕ числовые артикулы (#24,89)
  - Эмодзи внутри текстов
  - Цены (числа в формате 165.000, "60 000 сум", "-30%")
  - Призывы-триггеры скидок ("Только до X", "АКЦИЯ", "Распродажа", и т.д.)

Оставляется:
  - Описания товаров и брендов
  - Артикулы (Серия В, A27, Z1, #24,89)
  - Имена брендов (Lashy, Barbara, Dlux, и т.д.)
  - Полезный контент

Запуск:
    python clean_lashop_docs.py
"""

import re
import os
import sys
from docx import Document

sys.stdout.reconfigure(encoding="utf-8")

# ──────────────────────────────────────────────────────────────────────────
# REGEX'ы — паттерны для распознавания мусора
# ──────────────────────────────────────────────────────────────────────────

# Эмодзи: широкие unicode-диапазоны + пиктограммы + flag emoji + ZWJ-последовательности
EMOJI_RE = re.compile(
    "["
    "\U0001F300-\U0001FAFF"  # symbols & pictographs (большой блок)
    "\U0001F600-\U0001F64F"  # emoticons
    "\U0001F680-\U0001F6FF"  # transport & map symbols
    "\U0001F1E0-\U0001F1FF"  # flags (regional indicators)
    "\U00002600-\U000027BF"  # misc symbols, dingbats (❤, ❄, ☀ и т.д.)
    "\U0001F900-\U0001F9FF"  # supplemental symbols & pictographs
    "\U0001FA70-\U0001FAFF"  # symbols & pictographs extended-A
    "‍"  # ZWJ — связка эмодзи (❤️‍🔥)
    "️"  # variation selector-16 (делает ❤ → ❤️)
    "❤"  # heart
    "]+",
    flags=re.UNICODE,
)

# Точные строки-мусор (после strip)
EXACT_JUNK_LINES = {
    "🇺🇿LASHop- ВСЕ для наращивания ресниц и оформления бровей",
    "Photo",
    "Video file",
    "Video message",
    "Voice message",
    "Sticker",
    "Animation",
    "GIF",
    "File",
    "Audio file",
    "Anonymous poll",
    "Previous messages",
    "Not included, change data exporting settings to download.",
}

# Дата отдельной строкой: "27 December 2024"
DATE_LINE_RE = re.compile(
    r"^\d{1,2}\s+(January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{4}$"
)

# Время отдельной строкой: "18:04"
TIME_LINE_RE = re.compile(r"^\d{1,2}:\d{2}$")

# Числовая дата: формат с точкой 19.11.2024 или 19.11.24 (с опц. временем),
# либо формат со слэшем 19/11/2024 (год обязательно 4 цифры — иначе ловим
# перечисления размеров товара "3/5/10ml" как ложные срабатывания).
NUMERIC_DATE_RE = re.compile(
    r"\b\d{1,2}\.\d{1,2}\.\d{2,4}(?:\s+\d{1,2}:\d{2}(?::\d{2})?)?\b"
    r"|\b\d{1,2}/\d{1,2}/\d{4}\b"
)

# Размер фото: "1080×1080, 83.6 KB"
PHOTO_SIZE_RE = re.compile(r"^\d+×\d+,\s*[\d.]+\s*(KB|MB)$")

# Длительность+размер видео: "00:05, 1.1 MB"
VIDEO_DUR_RE = re.compile(r"^\d{1,2}:\d{2},\s*[\d.]+\s*(KB|MB)$")

# Реакция отдельной строкой: "❤4 👍2 🔥1" — эмодзи + цифра, разделено пробелами/nbsp
REACTION_LINE_RE = re.compile(
    r"^(?:[\U00002600-\U000027BF\U0001F300-\U0001FAFF‍️❤️]+\s*\d+[\s\xa0]*)+$"
)

# Ссылки внутри текста — удаляем полностью
URL_RE = re.compile(
    r"https?://\S+|t\.me/\S+|instagram\.com/\S+|www\.\S+",
    flags=re.IGNORECASE,
)

# Хэштеги-слова: #ресницы, #lash, #bzbridge — удаляем.
# Но #24,89 (числовой артикул) — НЕ удаляем.
HASHTAG_WORD_RE = re.compile(r"#[A-Za-zА-Яа-яёЁ_]+\w*")

# Юниты измерения для distinguishable формул акций ("10 мл по цене 5 мл")
UNIT = r"(?:мл|шт|г|гр|мг|кг|см|мм|м|л|пар|пара|пары|упак|упаковк[аиу])"

# Цены: 165.000, 115.500, 60.000, 84.000, 50 000 сум, 1 200 000 сум,
# а также диапазоны 100-150 (если рядом сум/тыс) — оставлю простой паттерн
PRICE_PATTERNS = [
    # 165.000, 115.500 (точка как разделитель тысяч) с опциональной валютой
    re.compile(r"\b\d{1,3}(?:\.\d{3})+(?:\s*(?:сум|сом|UZS|руб))?\b", flags=re.IGNORECASE),
    # 50 000 сум, 1 200 000 сум (пробел как разделитель)
    re.compile(r"\b\d{1,3}(?:[\s\xa0]\d{3})+(?:\s*(?:сум|сом|UZS|руб))?\b", flags=re.IGNORECASE),
    # 60000 сум — слитное число с валютой
    re.compile(r"\b\d{4,7}\s*(?:сум|сом|UZS|руб)\b", flags=re.IGNORECASE),
    # Скидки в процентах
    re.compile(r"-\d{1,3}\s*%"),
    # Формула акции "10 мл по цене 5 мл", "2 шт по цене 1 шт", "10мл — по цене 5мл"
    # (юниты с обеих сторон одинаковые) — удаляем фрагмент целиком вместе с числами и юнитами
    re.compile(
        rf"\b\d+\s*{UNIT}\s*[—–-]?\s*по\s+цене\s+\d+\s*{UNIT}\b",
        flags=re.IGNORECASE,
    ),
    # Цена без скидки / Цена со скидкой / стоимость / по цене
    re.compile(r"(?:Цена\s+(?:без\s+скидки|со\s+скидкой)|стоимость|по\s+цене|обычная\s+цена)\s*:?\s*[\d\s.]+", flags=re.IGNORECASE),
]

# Слова-триггеры акций — пользователь просил оставлять описания, поэтому
# просто удаляем эти слова, не весь пост
PROMO_WORD_RE = re.compile(
    r"\b(?:АКЦИЯ|акция|СКИДКА|скидка|РАСПРОДАЖА|распродажа|Ликвидация|"
    r"спецпредложение|только\s+до\s+\d+\s+\w+|осталось\s+\d+\s+(?:дня|дней|день)|"
    r"финальный\s+шанс|не\s+упустите|торопитесь|бери\s+больше\s+—\s+плати\s+меньше|"
    r"количество\s+ограничено|срок\s+годности[^.]*)\b",
    flags=re.IGNORECASE,
)

# Дата привязки внутри текста: "до 7 апреля", "до 7 марта"
DATE_REF_RE = re.compile(
    r"\bдо\s+\d{1,2}\s+(?:января|февраля|марта|апреля|мая|июня|июля|августа|сентября|октября|ноября|декабря)\b",
    flags=re.IGNORECASE,
)

# ──────────────────────────────────────────────────────────────────────────
# Логика очистки
# ──────────────────────────────────────────────────────────────────────────


def is_full_junk_line(line: str) -> bool:
    """Возвращает True если параграф целиком — системный мусор и должен быть удалён."""
    s = line.strip()
    if not s:
        return True
    if s in EXACT_JUNK_LINES:
        return True
    if DATE_LINE_RE.match(s):
        return True
    if TIME_LINE_RE.match(s):
        return True
    if NUMERIC_DATE_RE.fullmatch(s):
        return True
    if PHOTO_SIZE_RE.match(s):
        return True
    if VIDEO_DUR_RE.match(s):
        return True
    if REACTION_LINE_RE.match(s):
        return True
    # URL отдельной строкой
    if URL_RE.fullmatch(s):
        return True
    return False


def clean_inline_text(text: str) -> str:
    """Чистит текст внутри полезного параграфа: убирает эмодзи, ссылки, хэштеги,
    цены, промо-слова. Сохраняет описание товара и артикулы."""
    s = text

    # 1. Ссылки сначала — пока строка целая
    s = URL_RE.sub("", s)

    # 2. Хэштеги-слова (НЕ цифровые артикулы)
    s = HASHTAG_WORD_RE.sub("", s)

    # 3. Цены и скидочные шаблоны
    for pat in PRICE_PATTERNS:
        s = pat.sub("", s)

    # 4. Промо-триггеры
    s = PROMO_WORD_RE.sub("", s)

    # 5. Дата-привязки внутри текста ("до 7 апреля")
    s = DATE_REF_RE.sub("", s)

    # 5b. Числовые даты внутри текста (19.11.2024, 19.11.2024 20:25:39)
    s = NUMERIC_DATE_RE.sub("", s)

    # 6. Эмодзи (последним — потому что некоторые regexes выше могли их использовать)
    s = EMOJI_RE.sub("", s)

    # 7. Чистим типичные «висящие» знаки после удалений
    # — дублированные юниты («мл мл» → «мл»), которые остались после промо-чистки
    s = re.sub(rf"\b({UNIT})\s+\1\b", r"\1", s, flags=re.IGNORECASE)
    # — двойные пробелы, пробелы перед знаками препинания, пустые скобки, висящие тире
    s = re.sub(r"\(\s*\)", "", s)  # пустые скобки
    s = re.sub(r"\s+([,.!?:;])", r"\1", s)  # пробел перед знаком препинания
    s = re.sub(r"\s+\n", "\n", s)  # пробелы перед переносом
    s = re.sub(r"[ \t\xa0]+", " ", s)  # серии пробелов → один
    s = re.sub(r"\n{3,}", "\n\n", s)  # >2 пустых строк → 2
    # — висящие двоеточия / тире в конце или сразу после удалений
    s = re.sub(r":\s*$", "", s, flags=re.MULTILINE)
    s = re.sub(r"^\s*[—–-]\s*$", "", s, flags=re.MULTILINE)

    return s.strip()


def clean_docx(input_path: str, output_path: str) -> dict:
    """Загружает docx, обрабатывает каждый параграф, сохраняет очищенный."""
    src = Document(input_path)
    dst = Document()

    stats = {
        "input_paragraphs": 0,
        "kept_paragraphs": 0,
        "dropped_full_junk": 0,
        "dropped_after_clean_empty": 0,
    }

    for p in src.paragraphs:
        text = p.text
        stats["input_paragraphs"] += 1

        # 1. Полностью мусорные строки выкидываем целиком
        if is_full_junk_line(text):
            stats["dropped_full_junk"] += 1
            continue

        # 2. Чистим инлайн
        cleaned = clean_inline_text(text)

        if not cleaned:
            stats["dropped_after_clean_empty"] += 1
            continue

        dst.add_paragraph(cleaned)
        stats["kept_paragraphs"] += 1

    dst.save(output_path)
    return stats


# ──────────────────────────────────────────────────────────────────────────
# Прогон
# ──────────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    here = os.path.dirname(os.path.abspath(__file__))
    files = ["🇺🇿LASHop1.docx", "🇺🇿LASHop2.docx", "🇺🇿LASHop3.docx"]

    for f in files:
        in_path = os.path.join(here, f)
        out_name = f.replace(".docx", "_clean.docx")
        out_path = os.path.join(here, out_name)

        print(f"\n=== {f} ===")
        stats = clean_docx(in_path, out_path)
        print(f"  Input paragraphs:   {stats['input_paragraphs']}")
        print(f"  Dropped full junk:  {stats['dropped_full_junk']}")
        print(f"  Dropped → empty:    {stats['dropped_after_clean_empty']}")
        print(f"  KEPT:               {stats['kept_paragraphs']}")
        print(f"  → saved: {out_name}")
