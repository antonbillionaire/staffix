import sys, io, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE = r"c:/Users/anton/Documents/GITHUB REPOSITORIUM/staffix/Шаблоны документов/Reels_БЗ"

# Order of files in the document
files = [
    ("README.txt", "ВВЕДЕНИЕ И ИНСТРУКЦИЯ ПО ИСПОЛЬЗОВАНИЮ"),
    ("Reel_2.1_Регистрация_и_настройка.txt", "РОЛИК 2.1 — Регистрация и настройка за 10 минут"),
    ("Reel_2.2_Запись_клиента_через_AI.txt", "РОЛИК 2.2 — Как клиент записывается через AI"),
    ("Reel_2.3_База_знаний.txt", "РОЛИК 2.3 — База знаний: загрузите и AI выучит"),
    ("Reel_2.4_CRM_клиенты_и_сегменты.txt", "РОЛИК 2.4 — CRM: клиенты и сегменты"),
    ("Reel_2.5_Напоминания_и_отзывы.txt", "РОЛИК 2.5 — Напоминания и отзывы"),
    ("Reel_2.6_Каталог_товаров_и_заказы.txt", "РОЛИК 2.6 — Каталог товаров и заказы"),
    ("Reel_2.7_Аналитика.txt", "РОЛИК 2.7 — Аналитика: цифры вашего бизнеса"),
    ("Reel_2.8_Подключение_Instagram_WhatsApp.txt", "РОЛИК 2.8 — Подключение Instagram и WhatsApp"),
]

doc = Document()

# Document settings
section = doc.sections[0]
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2)

# Default font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("STAFFIX — Базы знаний для Instagram Reels")
run.bold = True
run.font.size = Pt(20)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = subtitle.add_run("Цикл 2: Обучение — 8 роликов")
sub_run.italic = True
sub_run.font.size = Pt(13)
sub_run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

doc.add_paragraph()

# Table of contents
toc_h = doc.add_paragraph()
toc_run = toc_h.add_run("Содержание")
toc_run.bold = True
toc_run.font.size = Pt(14)

for i, (_, title_text) in enumerate(files):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.add_run(f"{i + 1}. {title_text}").font.size = Pt(11)

doc.add_page_break()

# Process each file
for idx, (filename, section_title) in enumerate(files):
    filepath = os.path.join(BASE, filename)
    if not os.path.exists(filepath):
        print(f"Skipping missing: {filename}")
        continue

    # Section heading
    h = doc.add_paragraph()
    run = h.add_run(section_title)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1f, 0x4e, 0x79)

    # Read file content
    with open(filepath, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    in_block = False
    skip_next = False

    for line in lines:
        if skip_next:
            skip_next = False
            continue

        # Decoration line: skip
        if line.strip().startswith('=') and len(line.strip()) > 20:
            continue

        # Subsection heading (UPPERCASE between decoration lines)
        if line.strip() and line.strip().isupper() and len(line.strip()) > 5 and len(line.strip()) < 100:
            # Check if it's a heading by content
            stripped = line.strip()
            if any(kw in stripped for kw in ['ИНСТРУКЦИЯ', 'ШАГ', 'ВАЖНО', 'УСЛУГИ', 'МАСТЕР', 'ДЕМО',
                                              'РАСПИСАНИЕ', 'СЛОТЫ', 'FAQ', 'ВОПРОС', 'ОТВЕТ',
                                              'КАК AI', 'ЧТО', 'ПРИМЕР', 'СЕГМЕНТ', 'СТАТИСТИК',
                                              'НАСТРАИВАТЬ', 'ШАБЛОН', 'ЦИФРЫ', 'ПОДКЛЮЧЕН',
                                              'ССЫЛК', 'УВЕДОМЛЕН', 'ИНСАЙТ', 'СПИСОК',
                                              'СОВЕТ', 'РЕЦИКЛ', 'РЕАЛЬНЫЕ', 'ОДНА', 'ТЕСТИР',
                                              'ЗАПИСЫВАЙТЕ', 'ВВЕДЕНИЕ', 'ПОШАГОВАЯ',
                                              'ИНСТРУКЦИЯ', 'ТРЕБОВАНИЯ', 'ШАГИ', 'ВРЕМЯ',
                                              'ОСНОВНЫЕ', 'ЧАСТЫЕ', 'КАТАЛОГ', 'БУКЕТЫ',
                                              'КОМПОЗИЦ', 'ОДИНОЧН', 'ДОПОЛНИТЕЛЬНО', 'ЗОНЫ',
                                              'УПРАВЛЕН', 'СПЕЦИАЛЬН', 'СРАВНЕН', 'ЗАГРУЗКА',
                                              'ВРЕМЯ', 'ПОПУЛЯРН', 'ВЫРУЧКА', 'ПЕРИОД',
                                              'ОБЩАЯ', 'КЛИЕНТ', 'РЕКОМЕНД', 'ПРАЙС',
                                              'ПАРИКМАХЕР', 'ОКРАШИВАН', 'УКЛАДКА', 'УХОД',
                                              'НОГТЕВ', 'ЕДИНЫЙ', 'ВАЖНЫЕ', 'ТЕМА',
                                              'ХРОНОМЕТР', 'ХУК', 'ЗАКРЫТИЕ', 'ЗАМЕТКИ',
                                              'ПРОДАЖИ', 'ПОДКЛЮЧЕНИЕ', 'INSTAGRAM',
                                              'WHATSAPP', 'TELEGRAM', 'FACEBOOK',
                                              'ЭКСПОРТ', 'ПОДДЕРЖИВАЕМ']):
                p = doc.add_paragraph()
                run = p.add_run(stripped)
                run.bold = True
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(0x2e, 0x74, 0xb5)
                continue

        # Normal paragraph
        if line.strip():
            p = doc.add_paragraph(line.rstrip())
            p.paragraph_format.space_after = Pt(2)
        else:
            doc.add_paragraph()

    # Page break between sections (except last)
    if idx < len(files) - 1:
        doc.add_page_break()

# Save
output_path = r"c:/Users/anton/Documents/GITHUB REPOSITORIUM/staffix/Шаблоны документов/Staffix_Reels_БЗ_Полный_документ.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
print(f"File size: {os.path.getsize(output_path)} bytes")
